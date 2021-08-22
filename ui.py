import sys
import os
import math
import datetime
import time
import copy
import uuid
import collections
import json
import shutil

import xlrd
from PIL import Image
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtGui import QColor, QPixmap
from PyQt5.QtCore import Qt, QDateTime
from PyQt5.QtWidgets import QMainWindow, QMessageBox, QTableWidgetItem, QCheckBox, QPushButton, QApplication, \
    QHeaderView, QSplashScreen, QLineEdit
import pyqtgraph as pg
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from interface import Ui_MainWindow
from crawl import trans_time, check_login
from myMessage import MyMessageBox
from MyPlotWidget import MyPlotWidget
from crawlThread import CrawlThread
from settings import Settings
from draw_win import DrawWindow
from satelliteData import SatelliteData
from myProgress import MyProgress

pg.setConfigOption('background', 'w')
pg.setConfigOption('foreground', 'k')


class UiTest(QMainWindow, Ui_MainWindow):

    def __init__(self, parent=None):
        super(UiTest, self).__init__(parent)
        self.region = pg.RectROI([0, 0], [0, 0], pen=pg.mkPen('g', width=1))  # 框选
        self.excel_data = []  # 所有excel爬取的数据, 修改记录在这里
        self.crawl_status = False
        self.manual_item = None  # 当前操作的数据
        self.select_indexs = []  # 自动剔野选择行数组
        self.l_plot_data = None  # 左侧绘图数据
        self.r_plot_data = None  # 右侧绘图数据
        self.undo_list = []  # undo队列
        self.undo_base_point_list = []  # 变化率剔野基准点undo队列
        self.setupUi(self)
        self.binding_signal()
        self.extar_control()
        self.create_dir(['tmp/image', 'tmp/data', 'tmp/cache', 'source'])
        self.config = Settings()
        self.init_style()
        self.progress = MyProgress()

    def binding_signal(self):
        self.upload_excel_btn.clicked.connect(self.choose_excel)
        self.crawl_btn.clicked.connect(self.crawl)
        self.reset_crawl_btn.clicked.connect(self.reset_crawl)
        self.report_docx_btn.clicked.connect(self.report_excel)
        self.open_sdat_btn.clicked.connect(self.reload_sdat)
        self.fileinfo_table_2.itemChanged.connect(self.table_update)
        self.manual_btn.clicked.connect(lambda: self.edit_model(77))  # 77: M
        self.rate_btn.clicked.connect(lambda: self.edit_model(82))  # 82: R
        self.delete_btn.clicked.connect(self.delete_data)
        self.undo_btn.clicked.connect(self.delete_undo)
        self.undo_base_btn.clicked.connect(self.undo_base_point)
        self.save_change_btn.clicked.connect(self.save_chaneg)
        self.select_all_btn.clicked.connect(self.select_all)
        self.auto_choice_btn.clicked.connect(self.auto_choice)
        self.auto_choices_cbbox.currentIndexChanged.connect(self.change_auto_choice_index)
        self.save_auto_choice_config_btn.clicked.connect(self.save_auto_choice_config)
        self.data_get_btn.clicked.connect(lambda: self.select_tab('data_get'))
        self.data_analysis_btn.clicked.connect(lambda: self.select_tab('data_analysis'))
        self.choice_btn.clicked.connect(lambda: self.select_tab('choice'))
        self.select_range_btn.clicked.connect(self.select_range)
        self.zoom_in_btn.clicked.connect(self.zoom_in)
        self.zoom_out_btn.clicked.connect(self.zoom_out)

    def init_style(self):
        logo = QtGui.QPixmap('source/logo.png')
        self.label_3.setPixmap(logo)
        self.label_3.setScaledContents(True)
        self.hidden_frame('data_get')
        self.password_edit.setEchoMode(QLineEdit.Password)

        # 加载默认设置
        login_config = self.config.get_login()
        if login_config['username']:
            self.username_edit.setText(login_config['username'])
            self.password_edit.setText(login_config['password'])
        auto_choices = self.config.get_auto_choice_list()
        if auto_choices:
            choice_items = [str([item + 1 for item in indexs]) for indexs in auto_choices]
            choice_items.insert(0, '[]')
            self.auto_choices_cbbox.addItems(choice_items)

        self.reset_crawl_btn.setHidden(True)

        self.undo_base_btn.setEnabled(False)
        self.undo_base_btn.setStyleSheet(
            'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')

    def hidden_frame(self, tab):
        style_1 = 'font: 75 12pt "微软雅黑";background-color: rgb(255, 255, 255);color:#455ab3;border-top-left-radius:15px;border-top-right-radius:15px;'
        style_2 = 'background-color: rgb(80, 103, 203);font: 75 12pt "微软雅黑";color: rgb(255, 255, 255);border-top-left-radius:15px;border-top-right-radius:15px;'
        if tab == "data_get":
            self.frame.setHidden(False)
            self.fileinfo_table.setHidden(False)
            self.frame_7.setHidden(False)
            self.label_2.setHidden(False)
            self.open_sdat_btn.setHidden(False)

            self.frame_2.setHidden(True)
            self.fileinfo_table_2.setHidden(True)

            self.frame_3.setHidden(True)
            self.l_widget.setHidden(True)
            self.r_widget.setHidden(True)
            self.frame_4.setHidden(True)
            self.frame_6.setHidden(True)

            self.data_get_btn.setStyleSheet(style_1)
            self.data_analysis_btn.setStyleSheet(style_2)
            self.choice_btn.setStyleSheet(style_2)
        elif tab == 'data_analysis':
            self.frame.setHidden(True)
            self.fileinfo_table.setHidden(True)
            self.frame_7.setHidden(True)
            self.label_2.setHidden(True)
            self.open_sdat_btn.setHidden(True)

            self.frame_2.setHidden(False)
            self.fileinfo_table_2.setHidden(False)

            self.frame_3.setHidden(True)
            self.l_widget.setHidden(True)
            self.r_widget.setHidden(True)
            self.frame_4.setHidden(True)
            self.frame_6.setHidden(True)

            self.data_get_btn.setStyleSheet(style_2)
            self.data_analysis_btn.setStyleSheet(style_1)
            self.choice_btn.setStyleSheet(style_2)

        elif tab == 'choice':
            if self.manual_item is None:
                message_box = MyMessageBox()
                message_box.setContent("暂无数据", "请先选择要填充的数据")
                message_box.exec_()
                return

            self.frame.setHidden(True)
            self.fileinfo_table.setHidden(True)
            self.frame_7.setHidden(True)
            self.label_2.setHidden(True)
            self.open_sdat_btn.setHidden(True)

            self.frame_2.setHidden(True)
            self.fileinfo_table_2.setHidden(True)

            self.frame_3.setHidden(False)
            self.l_widget.setHidden(False)
            self.r_widget.setHidden(False)
            self.frame_4.setHidden(False)
            self.frame_6.setHidden(False)

            self.data_get_btn.setStyleSheet(style_2)
            self.data_analysis_btn.setStyleSheet(style_2)
            self.choice_btn.setStyleSheet(style_1)

    def create_dir(slef, path_list):
        for path in path_list:
            if not os.path.exists(path):
                os.makedirs(path)

    def reset_data(self):
        # 重置软件数据状态
        self.excel_data = []  # 所有excel爬取的数据, 修改记录在这里
        self.manual_item = None  # 当前操作的数据
        self.select_indexs = []  # 自动剔野选择行数组
        self.undo_list = []  # undo列表
        if self.l_plot_data:
            self.l_plot_data.setData(x=[], y=[], pen=pg.mkPen('g', width=1))
        if self.r_plot_data:
            self.r_plot_data.setData(x=[], y=[], pen=pg.mkPen('r', width=1))
        self.l_pw.reset_rate_edit()

    def extar_control(self):
        # 左边框
        l_plot_layout = QtWidgets.QGridLayout()  # 实例化一个网格布局层
        self.l_widget.setLayout(l_plot_layout)  # 设置K线图部件的布局层
        l_date_axis = TimeAxisItem(orientation='bottom')
        self.l_pw = MyPlotWidget(self, axisItems={'bottom': l_date_axis})  # 创建一个绘图控件
        self.l_pw.plotItem.setMouseEnabled(x=False)
        self.l_pw.showGrid(x=True, y=True)
        # 要将pyqtgraph的图形添加到pyqt5的部件中，我们首先要做的就是将pyqtgraph的绘图方式由window改为widget。PlotWidget方法就是通过widget方法进行绘图的
        self.l_widget.layout().addWidget(self.l_pw)

        # 右边框
        r_plot_layout = QtWidgets.QGridLayout()  # 实例化一个网格布局层
        self.r_widget.setLayout(r_plot_layout)  # 设置K线图部件的布局层
        r_date_axis = TimeAxisItem(orientation='bottom')
        self.r_pw = MyPlotWidget(self, axisItems={'bottom': r_date_axis})  # 创建一个绘图控件
        self.r_pw.plotItem.setMouseEnabled(x=False)
        self.r_pw.showGrid(x=True, y=True)
        # 要将pyqtgraph的图形添加到pyqt5的部件中，我们首先要做的就是将pyqtgraph的绘图方式由window改为widget。PlotWidget方法就是通过widget方法进行绘图的
        self.r_widget.layout().addWidget(self.r_pw)
        self.r_pw.region = self.region
        self.r_pw.l_widget = self.l_widget
        self.r_pw.position_lable = self.position_lable
        vLine = pg.InfiniteLine(angle=90, movable=False)
        hLine = pg.InfiniteLine(angle=0, movable=False)
        self.r_pw.addItem(vLine, ignoreBounds=True)
        self.r_pw.addItem(hLine, ignoreBounds=True)
        self.r_pw.vLine = vLine
        self.r_pw.hLine = hLine
        self.r_pw.undo_base_point_list = self.undo_base_point_list

    def resizeEvent(self, *args, **kwargs):
        self.fileinfo_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.fileinfo_table_2.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

    def choose_excel(self):
        if self.crawl_status:
            message_box = MyMessageBox()
            message_box.setContent("请等待", "请等待数据读取完成")
            message_box.exec_()
            return
        fileName_choose, filetype = QtWidgets.QFileDialog.getOpenFileName(self,
                                                                          "选取文件",
                                                                          os.getcwd(),  # 起始路径
                                                                          "Execl Files (*.xlsx;*.xls)")  # 设置文件扩展名过滤,用双分号间隔

        if fileName_choose == "":
            return
        # 重置软件状态
        self.reset_data()

        self.excel_path_edit.setText(fileName_choose)
        # 解析excel文件，填充内容
        self.parse_excel(fileName_choose)

    # 解析excel
    def parse_excel(self, file_path):
        # 清除原有数据
        self.excel_data = []

        # 打开上传 excel 表格
        file = xlrd.open_workbook(file_path)
        # 打开文件
        sheet_1 = file.sheet_by_index(0)  # 根据sheet页的排序选取sheet
        row_content = sheet_1.row_values(0)  # 获取指定行的数据，返回列表，排序自0开始
        row_number = sheet_1.nrows  # 获取有数据的最大行
        for i in range(1, row_number):
            telemetry_name = sheet_1.cell(i, 0).value
            telemetry_num = sheet_1.cell(i, 1).value
            normal_range = sheet_1.cell(i, 2).value
            telemetry_source = sheet_1.cell(i, 3).value
            img_num = sheet_1.cell(i, 4).value
            table_num = sheet_1.cell(i, 5).value
            params_one = sheet_1.cell(i, 6).value
            params_two = sheet_1.cell(i, 7).value
            params_three = sheet_1.cell(i, 8).value
            params_four = sheet_1.cell(i, 9).value
            create_time_text = self.create_time_edit.text()
            end_time_text = self.end_time_edit.text()
            start_time = self.trans_data_time(create_time_text)
            end_time = self.trans_data_time(end_time_text)
            dataHead = {
                "status": '未读取',
                'telemetry_name': telemetry_name,
                'telemetry_num': telemetry_num,
                'normal_range': normal_range,
                'telemetry_source': telemetry_source,
                'img_num': img_num,
                'table_num': table_num,
                'params_one': params_one,
                'params_two': params_two,
                'params_three': params_three,
                "params_four": params_four,
                'start_time': start_time,
                'end_time': end_time
            }
            file_path = 'tmp/data/' + telemetry_num + '_' + trans_time(create_time_text)[:-5] + '-' + trans_time(
                end_time_text)[:-5] + '.tmp'
            satellite_data = SatelliteData(dataHead=dataHead, file_path=file_path)
            self.excel_data.append(satellite_data)
        self.update_talbe1()
        self.crawl_btn.setEnabled(True)
        self.crawl_btn.setStyleSheet('font: 10pt "Microsoft YaHei UI";background-color:#455ab3;color:#fff;')

    def reload_sdat(self):
        # 手动加载数据
        if self.crawl_status:
            message_box = MyMessageBox()
            message_box.setContent("请等待", "请等待数据读取完成")
            message_box.exec_()
            return
        filename_list, filetype = QtWidgets.QFileDialog.getOpenFileNames(self,
                                                                         "选取文件",
                                                                         os.getcwd(),  # 起始路径
                                                                         "sDat Files (*.sDat)")
        if filename_list == []:
            return

        # 判断遥测代号是否重复
        old_telemetry_num_list = [satellite_data.dataHead['telemetry_num'] for satellite_data in self.excel_data]
        new_telemetry_num_list = []
        for filename in filename_list:
            tmp_star = SatelliteData(file_path=filename, dataHead=None)
            if tmp_star.dataHead['telemetry_num'] not in new_telemetry_num_list:
                new_telemetry_num_list.append(tmp_star.dataHead['telemetry_num'])

        error_num_list = []
        for item in new_telemetry_num_list:
            if item in old_telemetry_num_list:
                error_num_list.append(item)

        error_num_list = list(set(error_num_list))
        if error_num_list:
            message_box = MyMessageBox()
            message_box.setContent("请检查数据", "遥测代号已存在 \n" + str(error_num_list))
            message_box.exec_()
            return

        for filename in filename_list:
            star = SatelliteData(file_path=filename, dataHead=None)
            self.excel_data.append(copy.deepcopy(star))

        self.update_talbe1()
        self.crawl_btn.setEnabled(True)
        self.crawl_btn.setStyleSheet('font: 10pt "Microsoft YaHei UI";background-color:#455ab3;color:#fff;')

    def update_talbe1(self):
        count = self.fileinfo_table.rowCount()
        for i in range(count):
            self.fileinfo_table.removeRow(count - i - 1)

        excel_data = self.excel_data
        row = len(excel_data)
        self.fileinfo_table.setRowCount(len(excel_data))
        for r in range(row):
            item = excel_data[r].dataHead
            item_status = item['status']
            if item_status is None:
                one_text = '读取失败'
            elif item_status == '读取成功':
                one_text = '读取成功'
            else:
                one_text = item_status
            self.fileinfo_table.setItem(r, 0, QTableWidgetItem(str(one_text)))
            self.fileinfo_table.setItem(r, 1, QTableWidgetItem(str(item['telemetry_name'])))
            self.fileinfo_table.setItem(r, 2, QTableWidgetItem(str(item['telemetry_num'])))
            self.fileinfo_table.setItem(r, 3, QTableWidgetItem(str(item['normal_range'])))
            self.fileinfo_table.setItem(r, 4, QTableWidgetItem(str(item['telemetry_source'])))
            self.fileinfo_table.setItem(r, 5, QTableWidgetItem(str(item['img_num'])))
            self.fileinfo_table.setItem(r, 6, QTableWidgetItem(str(int(item['table_num']))))
            self.fileinfo_table.setItem(r, 7, QTableWidgetItem(str(item['params_one'])))
            self.fileinfo_table.setItem(r, 8, QTableWidgetItem(str(item['params_two'])))
            self.fileinfo_table.setItem(r, 9, QTableWidgetItem(str(item['params_three'])))
            self.fileinfo_table.setItem(r, 10, QTableWidgetItem(str(item['params_four'])))

            if item_status is None:
                self.fileinfo_table.item(r, 0).setBackground(QColor(255, 185, 15))
            elif item_status == '读取成功':
                self.fileinfo_table.item(r, 0).setBackground(QColor(100, 255, 0))

        # 表格1禁止编辑
        self.fileinfo_table.setEditTriggers(QtGui.QAbstractItemView.NoEditTriggers)

    def getSelectButton(self, id):
        # 手动剔野
        select_check_btn = QCheckBox()
        select_check_btn.setObjectName(str(id) + '_select')
        select_check_btn.setStyleSheet(''' text-align : center; margin : 40px; font : 13px  ''')
        select_check_btn.clicked.connect(lambda: self.select_row(id))
        return select_check_btn

    # 列表内添加按钮
    def buttonForRow(self, id):
        # 手动剔野
        updateBtn = QPushButton('手动剔野')
        updateBtn.setStyleSheet(''' text-align : center; background-color : Orange; margin : 5px;
                                    border-style: outmanual_choset; font: 10pt "Microsoft YaHei UI";  ''')
        updateBtn.clicked.connect(lambda: self.manual_choice(id))
        return updateBtn

    def update_talbe2(self):
        count = self.fileinfo_table_2.rowCount()
        for i in range(count):
            self.fileinfo_table_2.removeRow(count - i - 1)

        excel_data = self.excel_data
        row = len(excel_data)
        self.fileinfo_table_2.setRowCount(len(excel_data))
        for r in range(row):
            item = excel_data[r].dataHead
            selectBtn = self.getSelectButton(r)

            # 读取失败的判断
            tmp_text = "未剔野" if item['status'] is not None else "数据读取失败"
            if tmp_text == '数据读取失败':
                self.fileinfo_table_2.setItem(r, 1, QTableWidgetItem(tmp_text))
                self.fileinfo_table_2.item(r, 1).setBackground(QColor(255, 185, 15))
            else:
                if item['table_num'] in ['1', 1, '1.0', 1.0, '2', 2, '2.0', 2.0]:
                    self.fileinfo_table_2.setCellWidget(r, 0, selectBtn)
                self.fileinfo_table_2.setItem(r, 1, QTableWidgetItem(tmp_text))
                updateBtn = self.buttonForRow(r)
                self.fileinfo_table_2.setCellWidget(r, 12, updateBtn)
            self.fileinfo_table_2.setItem(r, 2, QTableWidgetItem(str(item['telemetry_name'])))
            self.fileinfo_table_2.setItem(r, 3, QTableWidgetItem(str(item['telemetry_num'])))
            self.fileinfo_table_2.setItem(r, 4, QTableWidgetItem(str(item['normal_range'])))
            self.fileinfo_table_2.setItem(r, 5, QTableWidgetItem(str(item['telemetry_source'])))
            self.fileinfo_table_2.setItem(r, 6, QTableWidgetItem(str(item['img_num'])))
            self.fileinfo_table_2.setItem(r, 7, QTableWidgetItem(str(int(item['table_num']))))
            self.fileinfo_table_2.setItem(r, 8, QTableWidgetItem(str(item['params_one'])))
            self.fileinfo_table_2.setItem(r, 9, QTableWidgetItem(str(item['params_two'])))
            self.fileinfo_table_2.setItem(r, 10, QTableWidgetItem(str(item['params_three'])))
            self.fileinfo_table_2.setItem(r, 11, QTableWidgetItem(str(item['params_four'])))


    def crawl_callback(self, msg):
        is_ok, content, satellite_data = msg
        index = self.excel_data.index(satellite_data)
        if is_ok:
            self.fileinfo_table.setItem(index, 0, QTableWidgetItem("读取成功"))
            self.fileinfo_table.item(index, 0).setBackground(QColor(100, 255, 0))
            p_v = self.progress.getValue() + (1 / len(self.excel_data)) * 100
            self.progress.setValue(p_v)
            self.progress.show()
            QApplication.processEvents()
        else:
            print('读取失败，telemetry_num = ', satellite_data.dataHead['telemetry_num'])
            satellite_data.dataHead['status'] = None
            self.reset_crawl_btn.setHidden(False)
            self.fileinfo_table.setItem(index, 0, QTableWidgetItem("读取失败"))
            self.fileinfo_table.item(index, 0).setBackground(QColor(255, 185, 15))

        if '未读取' not in [item.dataHead['status'] for item in self.excel_data]:
            self.crawl_status = False
            # 切换到数据分析标签，并填充数据
            self.hidden_frame('data_analysis')
            self.update_talbe2()
            self.progress.setValue(100)
            self.progress.hide()
            QApplication.processEvents()

    def crawl(self):
        if not self.excel_data:
            message_box = MyMessageBox()
            message_box.setContent("参数缺失", "未导入任何读取参数")
            message_box.exec_()
            return

        if self.crawl_status:
            message_box = MyMessageBox()
            message_box.setContent("请等待", "请等待数据读取完成")
            message_box.exec_()
            return

        # 判断遥测代号是否重复
        telemetry_num_list = [satellite_data.dataHead['telemetry_num'] for satellite_data in self.excel_data]
        error_num_list = []
        for telemetry_num in telemetry_num_list:
            if telemetry_num_list.count(telemetry_num) > 1 and telemetry_num not in error_num_list:
                error_num_list.append(telemetry_num)

        if error_num_list:
            message_box = MyMessageBox()
            message_box.setContent("请检查数据", "遥测代号发生重复 \n" + str(error_num_list))
            message_box.exec_()
            return

        # 判断是否存在需要爬取的条目
        if not [star for star in self.excel_data if (star.dataHead['status'] == '未读取' or star.dataHead['status'] == None)]:
            self.crawl_status = False
            # 切换到数据分析标签，并填充数据
            self.hidden_frame('data_analysis')
            self.update_talbe2()
            return

        # 检查各个控件参数
        username = self.username_edit.text()
        password = self.password_edit.text()
        model = self.model_edit.text()
        create_time = self.create_time_edit.text()
        end_time = self.end_time_edit.text()
        if username == '' or password == '' or model == '' or create_time == '' or end_time == '' or self.excel_data == []:
            message_box = MyMessageBox()
            message_box.setContent("参数缺失", "请完善参数信息")
            message_box.exec_()
            return

        # todo: 发布前记得复原
        # 判断账号密码是否正确
        try:
            is_login = check_login(username, password)
        except Exception as e:
            print('错误内容', e)
            message_box = MyMessageBox()
            message_box.setContent("登录失败", "网络连接失败")
            message_box.exec_()
            return

        if not is_login:
            message_box = MyMessageBox()
            message_box.setContent("读取失败", "账号或密码错误")
            message_box.exec_()
            return
        else:
            # 保存账户和密码
            self.config.change_login(self.username_edit.text(), self.password_edit.text())

        self.crawl_status = True
        self.config.change_login(self.username_edit.text(), self.password_edit.text())
        # 多线程爬取
        # 创建线程
        self.crawl_btn.setEnabled(False)
        self.crawl_btn.setStyleSheet('font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
        self.thread_list = []
        for satellite_data in self.excel_data:
            if satellite_data.dataHead['status'] == '未读取' or satellite_data.dataHead['status'] is None:
                # 更新可能修改的起止时间
                create_time_1 = self.trans_data_time(create_time)
                end_time_1 = self.trans_data_time(end_time)
                file_path = 'tmp/data/' + satellite_data.dataHead['telemetry_num'] + '_' + trans_time(create_time)[:-5] + '-' + trans_time(
                    end_time)[:-5] + '.tmp'
                satellite_data.dataHead['start_time'] = create_time_1
                satellite_data.dataHead['end_time'] = end_time_1
                satellite_data.dataHead['status'] = '未读取'
                satellite_data.file_path = file_path

                tmp_thread = CrawlThread(satellite_data, username, password, model, create_time, end_time)
                tmp_thread._signal.connect(self.crawl_callback)
                self.thread_list.append(tmp_thread)

        # 开始线程
        if self.thread_list:
            for thread in self.thread_list:
                thread.start()
                self.progress.setContent('读取中', '数据读取中，请等待')
                self.progress.setValue(5)
                self.progress.show()
                QApplication.processEvents()
        else:
            self.crawl_status = False
            # 切换到数据分析标签，并填充数据
            self.hidden_frame('data_analysis')
            self.update_talbe2()


    def reset_crawl(self):
        if self.crawl_status:
            message_box = MyMessageBox()
            message_box.setContent("请等待", "请等待数据读取完成")
            message_box.exec_()
            return

        # 重新下载错误的数据
        message_box = MyMessageBox()
        message_box.setContent("重新下载", "是否重新获取失败的数据")
        message_box.exec_()
        if message_box.reply == QMessageBox.Ok:
            self.crawl()
        else:
            return




    def manual_choice(self, r):
        '''
        手动剔野按钮事件
        :param r:
        :return:
        '''

        start = time.time()

        self.fileinfo_table_2.selectRow(r)
        self.update_choice_parms()
        self.manual_item = self.excel_data[r]
        satellite_data = self.excel_data[r]
        if not os.path.exists(satellite_data.file_path):
            message_box = MyMessageBox()
            message_box.setContent("获取失败", "数据文件位置可能发生变化")
            message_box.exec_()
            return

        satellite_data.resampling(satellite_data.dataHead['start_time'], satellite_data.dataHead['end_time'],
                                  self.progress)
        if self.manual_item.star_data == [] or self.manual_item.star_data is None:
            message_box = MyMessageBox()
            message_box.setContent("获取失败", "请检查数据后重新读取")
            message_box.exec_()
            return

        # 传递到手动剔野页面,更新数据
        l_x, l_y = self.get_choice_data_xy(self.manual_item.star_data)
        # 初次加载，采用同样的数据
        r_x = l_x
        r_y = l_y
        if self.l_plot_data == None:
            self.l_plot_data = self.l_pw.plot(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))  # 在绘图控件中绘制图形
            self.r_plot_data = self.r_pw.plot(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))
        else:
            self.l_plot_data.setData(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))
            self.r_plot_data.setData(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))

        # 手动剔野界面状态重置
        self.hidden_frame('choice')
        self.r_pw.reset_rate_edit()
        self.r_pw.is_manual_edit = False
        self.r_pw.is_rate_edit = False
        self.manual_lable.setText(self.manual_item.dataHead['telemetry_num'] + '-手动剔野')

        # 放大与缩小
        self.zoom_in_btn.setEnabled(False)
        self.zoom_in_btn.setStyleSheet(
            'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
        self.zoom_out_btn.setEnabled(False)
        self.zoom_out_btn.setStyleSheet(
            'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')

        self.manual_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
        self.rate_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')

        self.undo_base_btn.setEnabled(False)
        self.undo_base_btn.setStyleSheet(
            'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
        self.delete_btn.setEnabled(False)
        self.delete_btn.setStyleSheet(
            'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
        self.undo_btn.setEnabled(False)
        self.undo_btn.setStyleSheet(
            'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
        # self.save_change_btn.setEnabled(False)
        # self.save_change_btn.setStyleSheet(
        #     'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')

        self.region.setSize([0, 0], [0, 0])
        self.r_pw.roi_range = None
        self.r_pw.autoRange()
        self.l_pw.autoRange()

        end = time.time()
        print('手动踢野加载耗时: ', end - start)

    def select_row(self, r):
        pass
        # if self.raw_data[r]['data'] == [] or self.raw_data[r]['data'] is None:
        #     message_box = MyMessageBox()
        #     message_box.setContent("获取失败", "请检查数据后重新读取")
        #     message_box.exec_()
        #     object_name = str(r) + '_select'
        #     checkbox = self.findChild(QCheckBox, object_name)
        #     checkbox.setChecked(False)
        #     self.select_indexs = self.get_select_indexs()

    def get_select_indexs(self):
        # 获取当前勾选状态，0：未选中，2 已选中
        select_indexs = []
        number = len(self.excel_data)
        number_list = [i for i in range(number) if self.excel_data[i].dataHead['status'] is not None and self.excel_data[i].dataHead['table_num']
                       in ['1', 1, '1.0', 1.0, '2', 2, '2.0', 2.0]]
        for i in number_list:
            object_name = str(i) + '_select'
            checkbox = self.findChild(QCheckBox, object_name)
            if checkbox.checkState() == 2:
                select_indexs.append(i)
        return select_indexs

    # 标签页切换
    def select_tab(self, tab_type):
        if self.crawl_status:
            message_box = MyMessageBox()
            message_box.setContent("请等待", "请等待数据读取完成")
            message_box.exec_()
            return
        self.hidden_frame(tab_type)

    def select_range(self):
        if self.r_pw.is_zoom_edit:
            self.r_pw.is_zoom_edit = False
            self.select_range_btn.setStyleSheet(
                'background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
            self.zoom_in_btn.setEnabled(False)
            self.zoom_in_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.zoom_out_btn.setEnabled(False)
            self.zoom_out_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.r_pw.removeItem(self.region)
            self.region.setSize([0, 0], [0, 0])
            self.r_pw.roi_range = None

            if self.r_pw.is_manual_edit:
                self.edit_model(Qt.Key_M)
                self.edit_model(Qt.Key_M)
            elif self.r_pw.is_rate_edit:
                self.edit_model(Qt.Key_R)
                self.edit_model(Qt.Key_R)
            else:
                self.manual_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.rate_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')

                self.undo_base_btn.setEnabled(False)
                self.undo_base_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                self.delete_btn.setEnabled(False)
                self.delete_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                self.undo_btn.setEnabled(False)
                self.undo_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                # self.save_change_btn.setEnabled(False)
                # self.save_change_btn.setStyleSheet(
                #     'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')

            self.manual_btn.setEnabled(True)
            self.rate_btn.setEnabled(True)
        else:
            self.r_pw.is_zoom_edit = True
            range = self.region
            self.select_range_btn.setStyleSheet(
                'background-color : LightCoral;color:#fff;font: 10pt "Microsoft YaHei UI";')
            self.zoom_in_btn.setEnabled(True)
            self.zoom_in_btn.setStyleSheet(
                'background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
            self.zoom_out_btn.setEnabled(True)
            self.zoom_out_btn.setStyleSheet(
                'background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
            self.r_pw.addItem(self.region, ignoreBounds=True)

            # 下面按钮栏全部禁用
            self.manual_btn.setEnabled(False)
            self.manual_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.rate_btn.setEnabled(False)
            self.rate_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.undo_btn.setEnabled(False)
            self.undo_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.undo_base_btn.setEnabled(False)
            self.undo_base_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.delete_btn.setEnabled(False)
            self.delete_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            # self.save_change_btn.setEnabled(False)
            # self.save_change_btn.setStyleSheet(
            #     'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')

            # 重置选择框
            self.r_pw.region.setSize([0, 0])

    def zoom_in(self):
        # 根据选择区域放大页面
        select_range = self.r_pw.roi_range
        if select_range is None:
            return

        left_time = self.timestamp2timestr(select_range[0])
        right_time = self.timestamp2timestr(select_range[2])

        if left_time < self.manual_item.dataHead['start_time']:
            left_time = self.manual_item.dataHead['start_time']
        if right_time > self.manual_item.dataHead['end_time']:
            right_time = self.manual_item.dataHead['end_time']

        # 重新抽样
        self.manual_item.resampling(left_time, right_time, self.progress)
        l_x, l_y = self.get_choice_data_xy(self.manual_item.star_data)
        if self.manual_item.cache_list:
            # 针对缓存文件采样
            self.manual_item.resampling(left_time, right_time, self.progress, self.manual_item.cache_list[-1])
            r_x, r_y = self.get_choice_data_xy(self.manual_item.star_cache_data)
        else:
            r_x = l_x
            r_y = l_y
        if self.l_plot_data == None:
            self.l_plot_data = self.l_pw.plot(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))  # 在绘图控件中绘制图形
            self.r_plot_data = self.r_pw.plot(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))
        else:
            self.l_plot_data.setData(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))
            self.r_plot_data.setData(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))

        # 重置状态
        self.region.setSize([0, 0], [0, 0])
        self.r_pw.roi_range = None
        self.r_pw.autoRange()
        self.l_pw.autoRange()

    def zoom_out(self):
        # 默认缩放倍率为1
        scale = 2
        select_range = self.r_pw.roi_range
        star_data_list = list(self.manual_item.star_data.keys())
        left_stamp = self.timestr2timestamp(star_data_list[0])
        right_stamp = self.timestr2timestamp(star_data_list[-1])
        distance = (right_stamp - left_stamp) * (scale / 2)
        # 缩放50%
        left_time = self.timestamp2timestr(left_stamp - distance)
        right_time = self.timestamp2timestr(right_stamp + distance)

        # if select_range is None:
        #     # 如果选择区域为0，中心为区间中心
        #     left_stamp = self.timestr2timestamp(star_data_list[0])
        #     right_stamp = self.timestr2timestamp(star_data_list[-1])
        #     distance = (right_stamp - left_stamp) * (scale / 2)
        #     # 缩放50%
        #     left_time = self.timestamp2timestr(left_stamp - distance)
        #     right_time = self.timestamp2timestr(right_stamp + distance)
        # else:
        #     # todo 可以考虑优化缩小逻辑
        #     # 如果选择区域，中心点为选择区域中心
        #     left_stamp = select_range[0]
        #     right_stamp = select_range[2]
        #     left_time = self.timestamp2timestr()
        #     right_time = self.timestamp2timestr()

        if left_time < self.manual_item.dataHead['start_time']:
            left_time = self.manual_item.dataHead['start_time']
        if right_time > self.manual_item.dataHead['end_time']:
            right_time = self.manual_item.dataHead['end_time']

        # 重新抽样
        self.manual_item.resampling(left_time, right_time, self.progress)
        l_x, l_y = self.get_choice_data_xy(self.manual_item.star_data)
        if self.manual_item.cache_list:
            # 针对缓存文件采样
            self.manual_item.resampling(left_time, right_time, self.progress, self.manual_item.cache_list[-1])
            r_x, r_y = self.get_choice_data_xy(self.manual_item.star_cache_data)
        else:
            r_x = l_x
            r_y = l_y
        if self.l_plot_data == None:
            self.l_plot_data = self.l_pw.plot(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))  # 在绘图控件中绘制图形
            self.r_plot_data = self.r_pw.plot(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))
        else:
            self.l_plot_data.setData(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))
            self.r_plot_data.setData(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))

        # 重置状态
        self.region.setSize([0, 0], [0, 0])
        self.r_pw.roi_range = None
        self.r_pw.autoRange()
        self.l_pw.autoRange()

    def edit_model(self, key):
        if key == Qt.Key_M:
            self.rate_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
            self.undo_base_btn.setEnabled(False)
            self.undo_base_btn.setStyleSheet(
                'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            self.undo_base_btn.setHidden(True)
            self.r_pw.is_rate_edit = False
            self.r_pw.reset_rate_edit()
            # 手动剔野修改按钮背景色,以及编辑状态
            if self.r_pw.is_manual_edit:
                self.r_pw.is_manual_edit = False
                self.manual_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.r_pw.removeItem(self.region)
                self.delete_btn.setEnabled(False)
                self.delete_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                self.undo_btn.setEnabled(False)
                self.undo_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                # self.save_change_btn.setEnabled(False)
                # self.save_change_btn.setStyleSheet(
                #     'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            else:
                self.r_pw.is_manual_edit = True
                self.manual_btn.setStyleSheet(
                    'background-color : LightCoral;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.r_pw.addItem(self.region, ignoreBounds=True)
                self.delete_btn.setEnabled(True)
                self.delete_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.delete_btn.setText("批量删除(D)")
                self.undo_btn.setEnabled(True)
                self.undo_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                # self.save_change_btn.setEnabled(True)
                # self.save_change_btn.setStyleSheet(
                #     'background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
        elif key == Qt.Key_R:
            self.manual_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
            self.r_pw.is_manual_edit = False
            self.r_pw.removeItem(self.region)
            if self.r_pw.is_rate_edit:
                self.r_pw.is_rate_edit = False
                self.rate_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.undo_base_btn.setEnabled(False)
                self.undo_base_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')

                self.delete_btn.setEnabled(False)
                self.delete_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                self.undo_btn.setEnabled(False)
                self.undo_btn.setStyleSheet(
                    'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
                # self.save_change_btn.setEnabled(False)
                # self.save_change_btn.setStyleSheet(
                #     'font: 10pt "Microsoft YaHei UI";background-color:rgb(156,156,156);;color:#fff;')
            else:
                self.r_pw.is_rate_edit = True
                self.rate_btn.setStyleSheet('background-color : LightCoral;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.undo_base_btn.setEnabled(True)
                self.undo_base_btn.setHidden(False)
                self.undo_base_btn.setStyleSheet('font: 10pt "Microsoft YaHei UI";background-color:#455ab3;color:#fff;')
                self.delete_btn.setEnabled(True)
                self.delete_btn.setText("变化率剔野(D)")
                self.delete_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                self.undo_btn.setEnabled(True)
                self.undo_btn.setStyleSheet('background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')
                # self.save_change_btn.setEnabled(True)
                # self.save_change_btn.setStyleSheet(
                #     'background-color:#455ab3;color:#fff;font: 10pt "Microsoft YaHei UI";')

    def undo_base_point(self):
        self.r_pw.undo_base_line()

    def keyPressEvent(self, *args, **kwargs):
        key = args[0].key()
        if key == Qt.Key_M or key == Qt.Key_R:
            self.edit_model(key)
        elif key == Qt.Key_D:
            self.delete_data()
        elif key == Qt.Key_S:
            self.select_range()

    def get_choice_data_xy(self, star_data):
        x = []
        y = []
        for k, v in star_data.items():
            x.append(self.timestr2timestamp(k))
            y.append(float(v))
        return x, y

    # 删除数据
    def delete_data(self):
        if self.r_pw.is_manual_edit:  # 手动剔野
            # 选定区域
            select_range = self.r_pw.roi_range
            if select_range is None:
                return

            left_time = self.timestamp2timestr(select_range[0])
            right_time = self.timestamp2timestr(select_range[2])
            if left_time < self.manual_item.dataHead['start_time']:
                left_time = self.manual_item.dataHead['start_time']
            if right_time > self.manual_item.dataHead['end_time']:
                right_time = self.manual_item.dataHead['end_time']

            self.manual_item.manual_choice(left_time, right_time, float(select_range[1]), float(select_range[3]), self.progress)

            # 对缓存数据重新采样
            star_data_list = list(self.manual_item.star_data.keys())
            left_stamp = self.timestr2timestamp(star_data_list[0])
            right_stamp = self.timestr2timestamp(star_data_list[-1])
            # 当前区间
            curr_left_time = self.timestamp2timestr(left_stamp)
            curr_right_time = self.timestamp2timestr(right_stamp)
            self.manual_item.resampling(curr_left_time, curr_right_time, self.progress, self.manual_item.cache_list[-1])
            x, y = self.get_choice_data_xy(self.manual_item.star_cache_data)
            self.r_plot_data.setData(x=x, y=y, pen=pg.mkPen('r', width=1))
        elif self.r_pw.is_rate_edit:  # 变化率剔野
            # 获取基准点
            base_point = self.r_pw.base_point_list
            if not base_point:
                message_box = MyMessageBox()
                message_box.setContent("提示", "前先选择基本点")
                message_box.exec_()
                return

            normal_rate = self.manual_item.dataHead['params_three']
            if normal_rate is '' or normal_rate is None:
                message_box = MyMessageBox()
                message_box.setContent("提示", "变化率剔野参数缺失或错误")
                message_box.exec_()
                return
            else:
                try:
                    normal_rate = float(normal_rate.replace(' ',''))
                except:
                    message_box = MyMessageBox()
                    message_box.setContent("提示", "变化率剔野参数不是数字")
                    message_box.exec_()
                    return

            isTrue, data = self.manual_item.rate_choice(base_point, normal_rate, self.progress)
            if not isTrue:
                message_box = MyMessageBox()
                message_box.setContent("提示", data)
                message_box.exec_()
                return

            # 对缓存数据重新采样
            star_data_list = list(self.manual_item.star_data.keys())
            left_stamp = self.timestr2timestamp(star_data_list[0])
            right_stamp = self.timestr2timestamp(star_data_list[-1])
            # 当前区间
            curr_left_time = self.timestamp2timestr(left_stamp)
            curr_right_time = self.timestamp2timestr(right_stamp)
            self.manual_item.resampling(curr_left_time, curr_right_time, self.progress, self.manual_item.cache_list[-1])
            x, y = self.get_choice_data_xy(self.manual_item.star_cache_data)
            self.r_plot_data.setData(x=x, y=y, pen=pg.mkPen('r', width=1))
            # 删除基准点
            for i in range(len(base_point)):
                self.undo_base_point()

        # 重置状态
        self.region.setSize([0, 0], [0, 0])
        self.r_pw.roi_range = None
        self.r_pw.autoRange()
        self.l_pw.autoRange()

    def delete_undo(self):
        if self.manual_item.cache_list:
            cache = self.manual_item.undo_cache()
            if cache is None:
                r_x, r_y = self.get_choice_data_xy(self.manual_item.star_data)
                self.r_plot_data.setData(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))
            # 重新抽样
            self.manual_item.resampling(cache['start_time'], cache['end_time'], self.progress)
            l_x, l_y = self.get_choice_data_xy(self.manual_item.star_data)
            self.manual_item.resampling(cache['start_time'], cache['end_time'], self.progress, cache)
            r_x, r_y = self.get_choice_data_xy(self.manual_item.star_cache_data)
            if self.l_plot_data == None:
                self.l_plot_data = self.l_pw.plot(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))  # 在绘图控件中绘制图形
                self.r_plot_data = self.r_pw.plot(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))
            else:
                self.l_plot_data.setData(x=l_x, y=l_y, pen=pg.mkPen('g', width=1))
                self.r_plot_data.setData(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))
            if os.path.exists(cache['file_path']):
                os.remove(cache['file_path'])
        else:
            r_x, r_y = self.get_choice_data_xy(self.manual_item.star_data)
            self.r_plot_data.setData(x=r_x, y=r_y, pen=pg.mkPen('r', width=1))

        # 重置状态
        self.region.setSize([0, 0], [0, 0])
        self.r_pw.roi_range = None
        self.r_pw.autoRange()
        self.l_pw.autoRange()


    def save_chaneg(self):
        message_box = MyMessageBox()
        message_box.setContent("提示", "确定修改数据")
        message_box.exec_()

        if message_box.reply == QMessageBox.Ok:
            # 拷贝替换文件
            if self.manual_item.cache_list:
                source_file = self.manual_item.cache_list[-1]['file_path']
                dest_file = self.manual_item.file_path
                shutil.copyfile(source_file, dest_file)
                # 清理缓存文件
                self.manual_item.clear_cache_file()
            else:
                pass

            index = self.fileinfo_table_2.currentIndex().row()
            self.fileinfo_table_2.setItem(index, 1, QTableWidgetItem("手动剔野"))
            self.fileinfo_table_2.item(index, 1).setBackground(QColor(100, 255, 0))
            self.r_pw.is_edit = False
            self.hidden_frame('data_analysis')
            QApplication.processEvents()

    def select_all(self):
        number = len(self.excel_data)
        number_list = [i for i in range(number) if self.excel_data[i].dataHead['status'] is not None and self.excel_data[i].dataHead['table_num'] in
                       ['1', 1, '1.0', 1.0, '2', 2, '2.0', 2.0]]
        if self.select_all_btn.text() == '全选':
            for i in number_list:
                # if self.raw_data[i]['data'] == [] or self.raw_data[i]['data'] is None:
                #     continue
                object_name = str(i) + '_select'
                checkbox = self.findChild(QCheckBox, object_name)
                checkbox.setChecked(True)
            self.select_all_btn.setText('取消全选')
        elif self.select_all_btn.text() == '取消全选':
            for i in number_list:
                object_name = str(i) + '_select'
                checkbox = self.findChild(QCheckBox, object_name)
                checkbox.setChecked(False)
            self.select_all_btn.setText('全选')
        self.select_indexs = self.get_select_indexs()
        QApplication.processEvents()

    def update_choice_parms(self):
        # 通过界面剔野参数跟新源数据
        for star in self.excel_data:
            item = star.dataHead
            index = self.excel_data.index(star)
            item['normal_range'] = self.fileinfo_table_2.item(int(index), 4).text()
            item['telemetry_source'] = self.fileinfo_table_2.item(int(index), 5).text()
            item['img_num'] = self.fileinfo_table_2.item(int(index), 6).text()
            item['table_num'] = self.fileinfo_table_2.item(int(index), 7).text()
            item['params_one'] = self.fileinfo_table_2.item(int(index), 8).text()
            item['params_two'] = self.fileinfo_table_2.item(int(index), 9).text()
            item['params_three'] = self.fileinfo_table_2.item(int(index), 10).text()
            item['params_four'] = self.fileinfo_table_2.item(int(index), 11).text()

    def auto_choice(self):
        self.select_indexs = self.get_select_indexs()
        if not self.select_indexs:
            message_box = MyMessageBox()
            message_box.setContent("自动剔野", "请勾选自动剔野项")
            message_box.exec_()
            return

        self.update_choice_parms()

        tmp_data = collections.OrderedDict()
        for i in range(len(self.excel_data)):
            if i in self.select_indexs:
                tmp_data[str(i)] = self.excel_data[i]

        # 源包剔野
        # source = {"source": {"index": "value"}}
        start = time.time()
        source_type = collections.OrderedDict()
        for index, star in tmp_data.items():
            dataHead = star.dataHead
            if not self.check_choice("source", dataHead["params_four"]):
                continue
            if dataHead["telemetry_source"] not in source_type.keys():
                source_type[dataHead["telemetry_source"]] = collections.OrderedDict()
            source_type[dataHead["telemetry_source"]][index] = star

        for type, star_list in source_type.items():
            self.source_choice(star_list)
        end = time.time()
        print('源包剔野耗时: ', end - start)

        # 阈值剔野
        start = time.time()
        for index, star in tmp_data.items():
            # 表格二不进行阈值剔野
            if self.check_choice('threshold', star.dataHead['params_four']) and self.excel_data[i].dataHead['table_num'] in ['1', 1, '1.0', 1.0]:
                self.threshold_choice(index, star)
        end = time.time()
        print('阈值剔野耗时: ', end - start)
        # 变化率剔野和手动剔野已合并

        # 修改剔野状态
        for index in tmp_data.keys():
            self.fileinfo_table_2.setItem(int(index), 1, QTableWidgetItem("自动剔野"))
            self.fileinfo_table_2.item(int(index), 1).setBackground(QColor(100, 255, 0))
            QApplication.processEvents()

        message_box = MyMessageBox()
        message_box.setContent("自动剔野", "自动剔野已完成")
        message_box.exec_()

    # 选择自动剔野保存项
    def change_auto_choice_index(self, curr_index=None):
        if not curr_index:
            curr_index = self.auto_choices_cbbox.currentIndex()
        number = len(self.excel_data)
        number_list = [i for i in range(number) if self.excel_data[i].dataHead['status'] is not None and self.excel_data[i].dataHead['table_num'] in
                       ['1', 1, '1.0', 1.0, '2', 2, '2.0', 2.0]]
        # 清除原有选择
        for i in number_list:
            object_name = str(i) + '_select'
            checkbox = self.findChild(QCheckBox, object_name)
            checkbox.setChecked(False)
        self.select_indexs = []

        choice_lists = self.config.get_auto_choice_list()
        choice_lists.insert(0, [])
        for item in choice_lists[curr_index]:
            if item in number_list:
                # if self.raw_data[item]['data'] == [] or self.raw_data[item]['data'] is None:
                #     continue
                object_name = str(item) + '_select'
                checkbox = self.findChild(QCheckBox, object_name)
                checkbox.setChecked(True)
        self.select_indexs = self.get_select_indexs()

        QApplication.processEvents()

    # 保存自动剔野选项
    def save_auto_choice_config(self):
        self.select_indexs = self.get_select_indexs()
        if self.select_indexs == [] or self.select_indexs == None:
            message_box = MyMessageBox()
            message_box.setContent("配置保存", "当前选择为空")
            message_box.exec_()
        else:
            self.config.change_auto_choice(self.select_indexs)
            message_box = MyMessageBox()
            message_box.setContent("配置保存", "保存成功")
            message_box.exec_()
            auto_choices = self.config.get_auto_choice_list()
            if auto_choices:
                choice_items = [str([item + 1 for item in indexs]) for indexs in auto_choices]
                choice_items.insert(0, '[]')
                self.auto_choices_cbbox.clear()
                self.auto_choices_cbbox.addItems(choice_items)
                # 重新选择自动剔野项
                self.change_auto_choice_index(1)


    def tmp_write(self, cache_size, star, source_f, out_f):
        source_f.readline()
        star.dataHead['params_one'] = '[-0.5, 0.5]'
        star.dataHead['params_two'] = '[-0.5, 0.5]'
        dataHead = star.dataHead
        head = str(dataHead['status']) + '||' + str(dataHead['telemetry_name']) + '||' + str(
            dataHead['telemetry_num']) + '||' + str(dataHead['normal_range']) + '||' + str(
            dataHead['telemetry_source']) + '||' + str(dataHead['img_num']) + '||' + str(
            dataHead['table_num']) + '||' + str(dataHead['params_one']) + '||' + str(
            dataHead['params_two']) + '||' + str(dataHead['params_three']) + '||' + str(
            dataHead['params_four']) + '||' + str(dataHead['start_time']) + '||' + str(
            dataHead['end_time']) + '\n'
        out_f.write(head)

        while 1:
            source_lines = source_f.readlines(cache_size)
            if not source_lines:
                source_f.close()
                out_f.close()
                break
            out_f.writelines(source_lines)





    def source_choice(self, star_list):
        # {"index": "star"}
        # 长度小于5，直接返回
        if len(star_list) < 4:
            return

        self.progress.setContent("进度", '---源包剔野中---')
        self.progress.setValue(0)
        self.progress.show()
        QApplication.processEvents()
        # todo 在这里已经耗费了10个工时，没能找到特别好的优化方式，耗时主要是在文件的读取和写入
        # 获取选择文件中，数据最多的行数作为循环次数. 同时创建文件
        point_total = 0
        source_f_list = []
        tmp_f_list = []
        threshold_list = []
        for index_str, star in star_list.items():
            source_f_list.append(open(star.file_path, 'r', encoding='gbk'))
            tmp_file_name = star.file_path[:-5] + '.tmp'
            tmp_f_list.append(open(tmp_file_name, 'w', encoding='gbk'))
            total = star.bufcount(star.file_path)
            if total > point_total:
                point_total = total
            params_one = star.dataHead['params_one']
            threshold = params_one.replace("[", '').replace("]", '').replace(' ', '').replace('，', ',')
            threshold = threshold.split(',')
            threshold_list.append(threshold)

        # 写入文件头部信息
        # 并且读取第一次数据
        time_line_list = []
        time_str_list = []
        time_value_list = []
        for index, source_f in enumerate(source_f_list):
            line = source_f.readline()
            tmp_f_list[index].write(list(star_list.values())[index].get_headline())
            time_line = source_f.readline()
            time_line_list.append(time_line)
            time_line = time_line.replace('\n', '')
            time_str, time_value = time_line.split('||')
            time_str_list.append(time_str)
            time_value_list.append(time_value)

        # 依次步进剔野
        for i in range(point_total * len(star_list)):
            if i % 10000 == 0:
                self.progress.setValue((i / point_total) * 100)
                self.progress.show()
                QApplication.processEvents()
            if not source_f_list:
                break
            delete_index = []
            # 如果长度小于4了，直接写入
            if len(source_f_list) < 4:
                for j in range(len(source_f_list)):
                    source_line = source_f_list[j].readline()
                    line = source_line.replace('\n', '')
                    if line == '' or line is None:
                        delete_index.append(j)
                    else:
                        tmp_f_list[j].write(source_line)
                        tmp_f_list[j].flush()
            else:
                # 判断当前最小值
                min_str = min(time_str_list)
                # 判断最小值的个数
                number = time_str_list.count(min_str)
                delete_index = []
                if number >= 4:  # 同一时间段大于4个
                    # 判单时间最小值对应的野点格式有
                    n = 0
                    for index, time_str in enumerate(time_str_list):
                        if time_str == min_str:
                            if float(time_value_list[index]) > float(threshold_list[index][1]) or float(time_value_list[index]) < float(threshold_list[index][0]):
                                n = n + 1

                    for index, time_str in enumerate(time_str_list):
                        if time_str == min_str:
                            if n < 4:  # 野点数小于4 直接写入
                                tmp_f_list[index].write(time_line_list[index])
                            # else:
                            #     print(time_str)

                            source_line = source_f_list[index].readline()
                            line = source_line.replace('\n', '')
                            if line == '' or line is None:
                                delete_index.append(index)
                            else:
                                time_str, value = line.split('||')
                                time_line_list[index] = source_line
                                time_str_list[index] = time_str
                                time_value_list[index] = value
                else:  # 最小值写入并步进
                    for index, time_str in enumerate(time_str_list):
                        if time_str == min_str:
                            tmp_f_list[index].write(time_line_list[index])
                            source_line = source_f_list[index].readline()
                            line = source_line.replace('\n', '')
                            if line == '' or line is None:
                                delete_index.append(index)
                            else:
                                time_str, value = line.split('||')
                                time_line_list[index] = source_line
                                time_str_list[index] = time_str
                                time_value_list[index] = value
            if delete_index:
                # 避免在循环中删除循环本身的对象
                # 准备对象
                delete_source_f_list = []
                delete_tmp_f_list = []
                delete_time_line_list = []
                delete_time_str_list = []
                delete_time_value_list = []
                delete_threshold_list = []
                for d in delete_index:
                    source_f_list[d].close()
                    delete_source_f_list.append(source_f_list[d])
                    tmp_f_list[d].close()
                    delete_tmp_f_list.append(tmp_f_list[d])
                    delete_time_line_list.append(time_line_list[d])
                    delete_time_str_list.append(time_str_list[d])
                    delete_time_value_list.append(time_value_list[d])
                    delete_threshold_list.append(threshold_list[d])
                # 删除
                for d_index in range(len(delete_index)):
                    source_f_list.remove(delete_source_f_list[d_index])
                    tmp_f_list.remove(delete_tmp_f_list[d_index])
                    time_line_list.remove(delete_time_line_list[d_index])
                    time_str_list.remove(delete_time_str_list[d_index])
                    time_value_list.remove(delete_time_value_list[d_index])
                    threshold_list.remove(delete_threshold_list[d_index])
                delete_index.clear()


        self.progress.hide()
        # 保证文件关闭
        for f in source_f_list:
            f.close()
        for f in tmp_f_list:
            f.close()

        # 修改拷贝文件
        for index_str, star in star_list.items():
            if os.path.isfile(star.file_path):
                os.remove(star.file_path)
            tmp_file_name = star.file_path[:-5] + '.tmp'
            os.rename(tmp_file_name, star.file_path)



    def threshold_choice(self, index, star):
        params_two = star.dataHead['params_two']
        threshold = params_two.replace("[", '').replace("]", '').replace(' ', '').replace('，', ',')
        threshold = threshold.split(',')

        # 文件保存交互问题
        tmp_file_name = star.file_path[:-5] + '.tmp'
        source_f = open(star.file_path, 'r', encoding='gbk')
        tmp_f = open(tmp_file_name, 'w', encoding='gbk')
        source_line = source_f.readline()  # 跳过head行
        tmp_f.write(star.get_headline())
        self.progress.setContent("进度", '---' + star.dataHead['telemetry_num'] + '阈值剔野中---')
        self.progress.setValue(0)
        self.progress.show()
        QApplication.processEvents()
        progress_index = 0
        progress_number = star.bufcount(star.file_path) - 1
        cache_lines = []  # 满10000行再开始写入，加快速度
        progress_index = 0
        while 1:
            tmp_source_lines = source_f.readlines(100 * 1024 * 1024)
            if not tmp_source_lines:
                break
            for source_line in tmp_source_lines:
                progress_index = progress_index + 1
                line = source_line.replace('\n', '')
                if line == '':
                    continue
                else:
                    time_str, v = line.split('||')
                    if float(v) < float(threshold[0]) or float(v) > float(threshold[1]):
                        continue
                    cache_lines.append(source_line)

                if len(cache_lines) > 100000:
                    tmp_f.writelines(cache_lines)
                    tmp_f.flush()
                    cache_lines.clear()
                    # del cache_lines
                    self.progress.setValue((progress_index / progress_number) * 100)
                    self.progress.show()
                    QApplication.processEvents()

        if cache_lines:
            tmp_f.writelines(cache_lines)
            tmp_f.flush()
            cache_lines.clear()
        source_f.close()
        tmp_f.close()
        self.progress.hide()
        # 替换源文件
        #  如果文件存在，则删除
        if os.path.isfile(star.file_path):
            os.remove(star.file_path)
        os.rename(tmp_file_name, star.file_path)



    def check_choice(self, module_name, params_four):
        params = params_four.replace("(", '').replace("（", '').replace(')', '').replace('）', ''). \
            replace(' ', '').replace('，', ',').replace('[', '').replace(']', '').replace('【', '').replace('】', '')
        params = params.split(',')
        if len(params) != 3:
            return False
        if module_name == "source" and params[0] in [1, '1']:
            return True
        elif module_name == 'threshold' and params[1] in [1, '1']:
            return True
        elif module_name == 'rate' and params[2] in [1, '1']:
            return True
        return False

    def report_excel(self):
        # 选择路径与文件名
        reportname = os.path.join(os.getcwd(), datetime.datetime.now().strftime('%Y%m%d_%H%M') + '.docx')
        fileName_choose, filetype = QtWidgets.QFileDialog.getSaveFileName(self,
                                                                          "导出文件",
                                                                          reportname,  # 起始路径
                                                                          "Execl Files (*.docx;)")  # 设置文件扩展名过滤,用双分号间隔
        if fileName_choose == "":
            return

        # 剔野后json数据导出
        # todo 注意：由于文件大小以及磁盘空间的考虑，暂时取消文件导出，剔野过后的文件将替换原文件
        # file_dir = os.path.dirname(fileName_choose)
        # filename = time.strftime("%Y%m%d_%H%M%S") + '.json'
        # json_filename = os.path.join(file_dir, filename)
        # self.report_data_json(json_filename)

        self.update_choice_parms()
        # 准备数据
        create_time = self.create_time_edit.text()
        end_time = self.end_time_edit.text()

        start_y, start_m, start_d = create_time.split(' ')[0].split('/')
        end_y, end_m, end_d = end_time.split(' ')[0].split('/')
        error_number, records = self.create_table1_data()
        if error_number == 0:
            error_text = '无异常现象。'
        else:
            error_text = '异常' + str(error_number) + '次。'
        # 生成报告
        document = Document()
        # 标题一
        document.styles['Normal'].font.name = u'微软雅黑'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        head_text = '1.XXXX卫星' + start_y + '年' + start_m + '月' + start_d + '日' + '至' + end_m + '月' + end_d + '在轨维护报告'
        heading = document.add_heading('', level=0).add_run(head_text)
        heading.font.name = u'微软雅黑'
        heading._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        # 标题一简介
        intr = ''.join(["        ", start_y, '年', start_m, '月', start_d, '日至', end_y, '年', end_m, '月', end_d, '日',
                        'XXXX卫星在轨运行状态正常，卫星运行在XXXX模式，所查询数据均在安全范围以内，', error_text,
                        '具体情况见表1 在轨遥测数据监视情况记录表，和表2在轨状态位变化情况记录表。'])
        document.add_paragraph(intr)

        # 添加表格一
        p2 = document.add_paragraph("表一 在轨遥测数据监视情况记录表")
        p2 = p2.paragraph_format
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # p2.page_break_before = True
        table = document.add_table(rows=1, cols=6, style='Table Grid')
        table.autofit = True
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '监视内容'
        hdr_cells[2].text = '遥测代号'
        hdr_cells[3].text = '在轨范围'
        hdr_cells[4].text = '正常范围'
        hdr_cells[5].text = '判读结果'
        # records = self.create_table1_data()
        for order_number, test_content, telemetry_num, on_normal_range, normal_range, result in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(order_number + 1)
            row_cells[1].text = test_content
            row_cells[2].text = telemetry_num
            row_cells[3].text = self.range_cut(on_normal_range)
            row_cells[4].text = self.range_cut(normal_range)
            row_cells[5].text = result

        # 表格二
        p3 = document.add_paragraph("表二 在轨状态位变化情况记录表")
        p3 = p3.paragraph_format
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # p3.page_break_before = True
        table1 = document.add_table(rows=1, cols=5, style='Table Grid')
        table1.autofit = True
        h_cells = table1.rows[0].cells
        h_cells[0].text = '序号'
        h_cells[1].text = '遥测参数'
        h_cells[2].text = '遥测代号'
        h_cells[3].text = '状态位'
        # h_cells[4].text = '状态变化'
        records1 = self.create_table2_data()
        for order_number, telemetry_params_num, params_num, status_bit in records1:
            h_cells = table1.add_row().cells
            h_cells[0].text = str(order_number + 1)
            h_cells[1].text = telemetry_params_num
            h_cells[2].text = params_num
            h_cells[3].text = status_bit
            # h_cells[4].text = state_change

        one_head = '1.1 XXXX卫星控制系统性能在轨状况'
        one_heading = document.add_heading('', level=1).add_run(one_head)
        one_heading.font.name = u'微软雅黑'
        one_heading._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 多组数据组合制图
        drafting_number = {}
        data_list = [item for item in self.excel_data if item.dataHead['status'] is not None]
        for star in data_list:
            img_num = str(star.dataHead['img_num'])
            if img_num == '' or img_num == 'None':
                continue
            if img_num not in drafting_number:
                drafting_number[img_num] = []
            drafting_number[img_num].append(star)

        # 制图名排血
        img_nums = [k for k in drafting_number.keys()]
        img_nums.sort()
        # 生成图片
        for img_num in img_nums:
            drafting_list = drafting_number[img_num]
            image_path, h = self.create_docx_image(drafting_list)
            document.add_picture(image_path, width=Inches(6), height=Inches(1.5 * h))
            table1_title = document.add_paragraph("图" + img_num)
            table1_title = table1_title.paragraph_format
            table1_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        two_head = '1.2 二浮陀螺在轨运行状况'
        two_heading = document.add_heading('', level=1).add_run(two_head)
        two_heading.font.name = u'微软雅黑'
        two_heading._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        thr_head = '总结'
        thr_heading = document.add_heading('', level=1).add_run(thr_head)
        thr_heading.font.name = u'微软雅黑'
        thr_heading._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        document.save(fileName_choose)

        message_box = MyMessageBox()
        message_box.setContent("生成表格", "word文档生成成功")
        message_box.exec_()

    def table_update(self):
        # 在剔野中再统一修改剔野参数
        pass

    def create_table1_data(self):
        data = []
        error_number = 0
        self.progress.setContent("进度", '---表格统计生成中---')
        self.progress.setValue(0)
        self.progress.show()
        QApplication.processEvents()

        data_list = []
        for star in self.excel_data:
            table_num = star.dataHead['table_num'].replace(' ', '')
            if star.dataHead['status'] is not None and table_num in ['1', 1, '1.0', 1.0]:
                data_list.append(star)
        for index, star in enumerate(data_list):
            self.progress.setValue((index + 1) / len(self.excel_data) * 100)
            self.progress.show()
            QApplication.processEvents()
            item = star.dataHead
            # if item['data'] == [] or item['data'] is None:
            #     continue
            if '允许' in item['normal_range']:
                continue
            parms_range = star.get_data_range()
            normal_split = item['normal_range'].replace('°', '').replace('[', '').replace(']', '').replace('~',
                                                                                                           ',').replace(
                ' ', '').split(',')
            normal_range = [float(normal_split[0]), float(normal_split[1])]

            if parms_range[0] >= normal_range[0] and parms_range[1] <= normal_range[1]:
                range_status = "正常"
            else:
                error_number = error_number + 1
                range_status = "异常"
            child = (
                index, item['telemetry_name'], item["telemetry_num"], str(parms_range), str(normal_range), range_status)
            data.append(child)
        self.progress.hide()
        return error_number, tuple(data)
        # data = (
        #     (1, '滚动角', 'RKSA1', '该变量的范围（min,max）', '', '正常/异常'),
        #     (2, '滚动角', 'RKSA1', '该变量的范围（min,max）', '', '正常/异常'),
        #     (3, '滚动角', 'RKSA1', '该变量的范围（min,max）', '', '正常/异常'),
        #     (4, '滚动角', 'RKSA1', '该变量的范围（min,max）', '', '正常/异常'),
        #     (5, '滚动角', 'RKSA1', '该变量的范围（min,max）', '', '正常/异常'),
        # )

    def create_table2_data(self):
        data = []
        data_list = []
        num = 0
        for star in self.excel_data:
            table_num = star.dataHead['table_num'].replace(' ', '')
            if star.dataHead['status'] is not None and table_num in ['2', 2, '2.0', 2.0]:
                data_list.append(star)
        for index, star in enumerate(data_list):
            # 取最后一个值作为状态
            item = star.dataHead
            last_data = star.get_last_data()
            child = (
                num, item['telemetry_name'], item["telemetry_num"], last_data)
            num = num + 1
            data.append(child)
        return tuple(data)

    def split_image(self, source_paths, save_path, flag='horizontal'):
        """
        :param source_paths: 原始图片数组
        :param save_path: 保存路径
        :param flag: horizontal or vertical
        :return:
        """
        # 计算图片长宽
        w_size = 0
        h_size = 0
        loc_list = []
        if flag == 'horizontal':
            for item in source_paths:
                loc_list.append((w_size, 0))
                img = Image.open(item)
                size = img.size
                w_size = w_size + size[0]
                h_size = size[1]
        elif flag == 'vertical':
            for item in source_paths:
                loc_list.append((0, h_size))
                img = Image.open(item)
                size = img.size
                w_size = size[0]
                h_size = h_size + size[1]

        # 拼接图片
        joint = Image.new('RGB', (w_size, h_size))
        for item in source_paths:
            img = Image.open(item)
            joint.paste(img, loc_list[source_paths.index(item)])

        joint.save(save_path)

    def create_docx_image(self, star_list):
        color_list = [(220, 20, 60), (0, 0, 255), (0, 255, 0), (255, 140, 0), (0, 255, 255), (255, 0, 255), (0, 0, 139)]
        index = 0
        split_name_list = []
        for star in star_list:
            color = color_list[index % len(color_list)]
            index = index + 1
            star.resampling(star.dataHead['start_time'], star.dataHead['end_time'],  self.progress)
            x, y = self.get_choice_data_xy(star.star_data)
            draw = DrawWindow(index, color, x, y, split_name_list)
            draw.showFullScreen()

        # 拼接图片
        file_name = 'tmp/image/' + str(uuid.uuid1()).replace('-', '') + '.png'
        self.split_image(split_name_list, file_name, flag='vertical')
        return file_name, len(split_name_list)

    def report_data_json(self, filename):
        # 返回剔野后的数据，暂定json格式
        data = self.excel_data
        with open(filename, 'w') as outfile:
            json.dump(data, outfile)

    def range_cut(self, data_str):
        rtn = str([format(float(item), '.4f') for item in json.loads(data_str)]).replace("'", '')
        return rtn



    def timestr2timestamp(self, timestr):
        datetime_obj = datetime.datetime.strptime(timestr, "%Y-%m-%d %H:%M:%S.%f")
        ret_stamp = int(time.mktime(datetime_obj.timetuple()) * 1000.0 + datetime_obj.microsecond / 1000.0)
        return ret_stamp / 1000

    def timestamp2timestr(self, timestamp):
        # 1602741526.7983296
        datetime_str = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(timestamp))
        return datetime_str + '.000000'

    def trans_data_time(self, time_str):
        '''
        将控件时间转换为和数据时间一样的格式
        :param time_str:
        :return:
        '''
        #  2020-10-11 18:25:27.454000
        l = time_str.split(" ")
        s = l[1].split(":")
        d = l[0].replace('/', '-')
        d = '-'.join([i if len(i) > 1 else '0' + i for i in d.split('-')])
        h = s[0] if len(s[0]) == 2 else '0' + s[0]
        m = s[1]
        s = '00000'
        new_time = d + ' ' + h + ':' + m + ':00.00000'
        return new_time


class TimeAxisItem(pg.AxisItem):
    def tickStrings(self, values, scale, spacing):
        return [datetime.datetime.fromtimestamp(value) for value in values]


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    splash = QSplashScreen(QPixmap(r"source/wait.png"))  # 启动界面图片地址
    splash.show()
    app.processEvents()
    ui = UiTest()
    ui.show()
    splash.finish(ui)
    sys.exit(app.exec_())
